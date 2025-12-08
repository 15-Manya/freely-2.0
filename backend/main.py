from contextlib import asynccontextmanager
from datetime import datetime
from typing import List, Optional
from fastapi import Depends, FastAPI, HTTPException, status, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from bson import ObjectId
import io
import csv
from docx import Document

from deps_auth import get_current_user
from database import connect_to_mongo, close_mongo_connection, get_database
from models import RiskAnalysis
from services.ai_service import analyze_client_chat, generate_proposal, update_proposal


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Manage application lifespan events."""
    # Startup: Connect to MongoDB
    await connect_to_mongo()
    yield
    # Shutdown: Close MongoDB connection
    await close_mongo_connection()


app = FastAPI(title="FreeLy API", version="0.1.0", lifespan=lifespan)

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # Frontend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/")
async def root():
    return {"message": "Welcome to FreeLy API"}


@app.get("/api/health")
async def health_check():
    """Health check endpoint."""
    from database import get_database
    
    db = get_database()
    if db is None:
        return {"status": "unhealthy", "database": "not_connected"}
    
    try:
        # Test database connection
        await db.command("ping")
        return {"status": "healthy", "database": "connected"}
    except Exception as e:
        return {"status": "unhealthy", "database": "error", "error": str(e)}


@app.get("/api/test-db")
async def test_database():
    """
    Test endpoint to verify MongoDB is working.
    Creates a test document, reads it, and deletes it.
    """
    from database import get_database
    from datetime import datetime
    
    db = get_database()
    if db is None:
        return {
            "success": False,
            "message": "Database not connected",
            "steps": {
                "connection": "failed",
                "write": "skipped",
                "read": "skipped",
                "delete": "skipped"
            }
        }
    
    test_collection = db["test"]
    test_id = f"test_{datetime.utcnow().timestamp()}"
    results = {
        "connection": "success",
        "write": None,
        "read": None,
        "delete": None
    }
    
    try:
        # Step 1: Test connection
        await db.command("ping")
        results["connection"] = "success"
        
        # Step 2: Write test document
        test_doc = {
            "_id": test_id,
            "message": "This is a test document",
            "timestamp": datetime.utcnow().isoformat(),
            "test": True
        }
        await test_collection.insert_one(test_doc)
        results["write"] = "success"
        
        # Step 3: Read test document
        found_doc = await test_collection.find_one({"_id": test_id})
        if found_doc:
            results["read"] = "success"
            results["read_data"] = {
                "id": str(found_doc.get("_id")),
                "message": found_doc.get("message"),
                "timestamp": found_doc.get("timestamp")
            }
        else:
            results["read"] = "failed - document not found"
        
        # Step 4: Delete test document
        delete_result = await test_collection.delete_one({"_id": test_id})
        if delete_result.deleted_count > 0:
            results["delete"] = "success"
        else:
            results["delete"] = "failed - document not deleted"
        
        # Get database stats
        stats = await db.command("dbStats")
        
        return {
            "success": True,
            "message": "Database test completed successfully",
            "database_name": db.name,
            "steps": results,
            "database_stats": {
                "collections": stats.get("collections", 0),
                "data_size": stats.get("dataSize", 0),
                "storage_size": stats.get("storageSize", 0)
            }
        }
        
    except Exception as e:
        return {
            "success": False,
            "message": f"Database test failed: {str(e)}",
            "steps": results,
            "error": str(e)
        }


@app.get("/api/me")
async def read_current_user(user=Depends(get_current_user)):
    """
    Return basic information about the authenticated user.

    Expects a valid Firebase ID token (JWT) in the Authorization header:
    Authorization: Bearer <id_token>
    """
    print(f"✅ Token verified successfully for user: {user.get('user_id')}")
    return {
        "uid": user.get("user_id"),
        "email": user.get("email"),
        "email_verified": user.get("email_verified", False),
        "name": user.get("name"),
        "picture": user.get("picture"),
    }


@app.post("/api/risk-analysis")
async def create_risk_analysis(
    analysis_type: str = Form(...),
    client_name: Optional[str] = Form(None),
    chat_file: Optional[UploadFile] = File(None),
    user=Depends(get_current_user)
):
    """
    Create a new risk analysis.
    For client_chat_import: requires client_name and chat_file (no PDFs).
    """
    db = get_database()
    if db is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Database not available"
        )
    
    user_id = user.get("user_id")
    
    # Validate analysis type
    valid_types = ["client_chat_import", "job_proposal", "text"]
    if analysis_type not in valid_types:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid analysis_type. Must be one of: {', '.join(valid_types)}"
        )
    
    # Validate client_chat_import requirements
    if analysis_type == "client_chat_import":
        if not client_name:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="client_name is required for client_chat_import"
            )
        if not chat_file:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="chat_file is required for client_chat_import"
            )
        # Check file type (no PDFs)
        if chat_file.filename:
            file_ext = chat_file.filename.split('.')[-1].lower()
            if file_ext == 'pdf':
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="PDF files are not supported. Please upload a text document (.txt, .docx, .csv)."
                )
            # Note: .doc (old Word format) is not fully supported
            if file_ext == 'doc':
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Old .doc files are not supported. Please convert to .docx or upload as .txt file."
                )
    
    # Read file content if provided
    file_content = None
    file_name = None
    file_size = 0
    text_content = None
    
    if chat_file:
        file_content = await chat_file.read()
        file_name = chat_file.filename
        file_size = len(file_content)
        
        # Determine file type and extract text accordingly
        if file_name:
            file_ext = file_name.split('.')[-1].lower()
            
            # Handle DOCX files (Microsoft Word)
            if file_ext == 'docx':
                try:
                    from docx import Document
                    from io import BytesIO
                    doc = Document(BytesIO(file_content))
                    # Extract all text from paragraphs
                    text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    # Also extract text from tables if any
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text_content += '\n' + cell.text
                except Exception as e:
                    print(f"⚠️  Error parsing DOCX file: {e}")
                    text_content = None
            
            # Handle DOC files (older Word format - requires additional library)
            elif file_ext == 'doc':
                # DOC files are binary and harder to parse
                # For now, we'll try to extract what we can or show an error
                text_content = None
                print("⚠️  .doc files are not fully supported. Please convert to .docx or .txt")
            
            # Handle plain text files
            elif file_ext in ['txt', 'csv']:
                try:
                    # Try UTF-8 first
                    text_content = file_content.decode('utf-8')
                except UnicodeDecodeError:
                    try:
                        # Try latin-1 as fallback
                        text_content = file_content.decode('latin-1')
                    except UnicodeDecodeError:
                        text_content = None
            
            # For other file types, try to decode as text
            else:
                try:
                    text_content = file_content.decode('utf-8')
                except UnicodeDecodeError:
                    try:
                        text_content = file_content.decode('latin-1')
                    except UnicodeDecodeError:
                        text_content = None
        
        # Store full file content for analysis
        # In production, consider storing large files in cloud storage (S3, GCS, etc.)
        input_data = {
            "file_name": file_name,
            "file_size": file_size,
            "file_type": chat_file.content_type or "unknown",
            "chat_content": text_content,  # Full chat content for analysis
            "has_full_content": text_content is not None
        }
    else:
        input_data = {}
    
    # Create risk analysis document
    risk_analysis_doc = {
        "user_id": user_id,
        "analysis_type": analysis_type,
        "client_name": client_name,
        "status": "pending",
        "input_data": input_data,
        "results": None,
        "created_at": datetime.utcnow(),
        "updated_at": datetime.utcnow()
    }
    
    # Insert into database
    collection = db["risk_analyses"]
    result = await collection.insert_one(risk_analysis_doc)
    analysis_id = result.inserted_id
    
    # Start analysis in background (for client_chat_import)
    if analysis_type == "client_chat_import" and text_content:
        # Update status to processing
        await collection.update_one(
            {"_id": analysis_id},
            {"$set": {"status": "processing", "updated_at": datetime.utcnow()}}
        )
        
        # Run analysis in background (using asyncio.create_task)
        import asyncio
        asyncio.create_task(process_risk_analysis(str(analysis_id), text_content, collection))
    elif analysis_type == "client_chat_import" and not text_content:
        # If no text content, mark as failed
        await collection.update_one(
            {"_id": analysis_id},
            {
                "$set": {
                    "status": "failed",
                    "results": {"error": "Failed to extract text from uploaded file"},
                    "updated_at": datetime.utcnow()
                }
            }
        )
    
    # Return created analysis
    created_doc = await collection.find_one({"_id": analysis_id})
    created_doc["_id"] = str(created_doc["_id"])
    
    return created_doc


async def process_risk_analysis(analysis_id: str, client_chat: str, collection):
    """
    Process risk analysis in the background.
    
    Args:
        analysis_id: MongoDB document ID as string
        client_chat: The client chat content
        collection: MongoDB collection reference
    """
    try:
        # Run AI analysis
        analysis_results = await analyze_client_chat(client_chat)
        
        # Update document with results
        await collection.update_one(
            {"_id": ObjectId(analysis_id)},
            {
                "$set": {
                    "status": "completed",
                    "results": analysis_results,
                    "updated_at": datetime.utcnow()
                }
            }
        )
        print(f"✅ Risk analysis completed for analysis_id: {analysis_id}")
        
    except Exception as e:
        import traceback
        error_msg = str(e)
        print(f"❌ Risk analysis failed for analysis_id: {analysis_id}: {error_msg}")
        traceback.print_exc()
        
        # Update document with error status
        try:
            await collection.update_one(
                {"_id": ObjectId(analysis_id)},
                {
                    "$set": {
                        "status": "failed",
                        "results": {"error": error_msg},
                        "updated_at": datetime.utcnow()
                    }
                }
            )
        except Exception as update_error:
            print(f"❌ Failed to update error status: {update_error}")


@app.get("/api/risk-analysis")
async def get_risk_analyses(
    user=Depends(get_current_user),
    limit: int = 50,
    skip: int = 0
):
    """Get all risk analyses for the current user."""
    db = get_database()
    if db is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Database not available"
        )
    
    user_id = user.get("user_id")
    collection = db["risk_analyses"]
    
    # Query user's analyses, sorted by created_at descending
    cursor = collection.find({"user_id": user_id}).sort("created_at", -1).skip(skip).limit(limit)
    analyses = await cursor.to_list(length=limit)
    
    # Convert ObjectId to string
    for analysis in analyses:
        analysis["_id"] = str(analysis["_id"])
    
    return {"analyses": analyses, "count": len(analyses)}


@app.get("/api/risk-analysis/{analysis_id}")
async def get_risk_analysis(
    analysis_id: str,
    user=Depends(get_current_user)
):
    """Get a specific risk analysis by ID."""
    db = get_database()
    if db is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Database not available"
        )
    
    user_id = user.get("user_id")
    collection = db["risk_analyses"]
    
    try:
        analysis = await collection.find_one({
            "_id": ObjectId(analysis_id),
            "user_id": user_id
        })
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid analysis ID"
        )
    
    if not analysis:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Risk analysis not found"
        )
    
    analysis["_id"] = str(analysis["_id"])
    return analysis


@app.delete("/api/risk-analysis/{analysis_id}")
async def delete_risk_analysis(
    analysis_id: str,
    user=Depends(get_current_user)
):
    """Delete a specific risk analysis by ID."""
    db = get_database()
    if db is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Database not available"
        )
    
    user_id = user.get("user_id")
    collection = db["risk_analyses"]
    
    try:
        # Find the analysis and verify ownership
        analysis = await collection.find_one({
            "_id": ObjectId(analysis_id),
            "user_id": user_id
        })
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid analysis ID"
        )
    
    if not analysis:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Risk analysis not found"
        )
    
    # Delete the analysis
    delete_result = await collection.delete_one({
        "_id": ObjectId(analysis_id),
        "user_id": user_id
    })
    
    if delete_result.deleted_count == 0:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Risk analysis not found or already deleted"
        )
    
    return {"message": "Risk analysis deleted successfully", "deleted_id": analysis_id}


@app.post("/api/risk-analysis/{analysis_id}/generate-proposal")
async def generate_proposal_from_risk_analysis(
    analysis_id: str,
    user=Depends(get_current_user)
):
    """
    Generate a proposal from an existing risk analysis.
    Uses the chat content from the risk analysis's input_data.
    """
    db = get_database()
    if db is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Database not available"
        )
    
    user_id = user.get("user_id")
    collection = db["risk_analyses"]
    
    # Find the risk analysis
    try:
        analysis = await collection.find_one({
            "_id": ObjectId(analysis_id),
            "user_id": user_id
        })
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid risk analysis ID"
        )
    
    if not analysis:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Risk analysis not found or not owned by user"
        )
    
    # Check if analysis has chat content
    if not analysis.get("input_data") or not analysis.get("input_data", {}).get("chat_content"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Risk analysis does not have chat content. Cannot generate proposal."
        )
    
    chat_content = analysis["input_data"]["chat_content"]
    client_name = analysis.get("client_name") or "Untitled Proposal"
    
    # Create proposal document
    proposal_collection = db["proposals"]
    proposal_doc = {
        "user_id": user_id,
        "proposal_type": "from_chat",
        "client_name": client_name,
        "status": "pending",
        "input_data": {
            "file_name": analysis["input_data"].get("file_name", "from_risk_analysis"),
            "file_size": analysis["input_data"].get("file_size", 0),
            "file_type": analysis["input_data"].get("file_type", "unknown"),
            "chat_content": chat_content,
            "has_full_content": True
        },
        "results": None,
        "created_at": datetime.utcnow(),
        "updated_at": datetime.utcnow()
    }
    
    # Insert into database
    result = await proposal_collection.insert_one(proposal_doc)
    proposal_id = result.inserted_id
    
    # Update status to processing
    await proposal_collection.update_one(
        {"_id": proposal_id},
        {"$set": {"status": "processing", "updated_at": datetime.utcnow()}}
    )
    
    # Start proposal generation in background
    import asyncio
    asyncio.create_task(process_proposal(str(proposal_id), chat_content, proposal_collection))
    
    # Return created proposal
    created_doc = await proposal_collection.find_one({"_id": proposal_id})
    created_doc["_id"] = str(created_doc["_id"])
    
    return created_doc


async def process_proposal(proposal_id: str, chat_content: str, collection):
    """
    Background task to process proposal generation using AI.
    
    Args:
        proposal_id: MongoDB document ID as string
        chat_content: The client chat text content
        collection: MongoDB collection for proposals
    """
    try:
        # Perform AI proposal generation
        ai_results = await generate_proposal(chat_content)
        
        # Initialize history with the first version
        first_version = {
            "formatted_proposal": ai_results.get("formatted_proposal", ""),
            "timestamp": datetime.utcnow(),
            "version": 1
        }
        
        # Update proposal status and results in DB
        await collection.update_one(
            {"_id": ObjectId(proposal_id)},
            {
                "$set": {
                    "status": "completed",
                    "results": ai_results,
                    "history": [first_version],
                    "current_version_index": 0,  # First version is at index 0
                    "updated_at": datetime.utcnow()
                }
            }
        )
        print(f"✅ Proposal {proposal_id} completed and updated in DB.")
    except Exception as e:
        error_msg = f"AI proposal generation failed: {str(e)}"
        print(f"❌ Error processing proposal {proposal_id}: {error_msg}")
        # Update status to failed
        try:
            await collection.update_one(
                {"_id": ObjectId(proposal_id)},
                {
                    "$set": {
                        "status": "failed",
                        "results": {"error": error_msg},
                        "updated_at": datetime.utcnow()
                    }
                }
            )
        except Exception as update_error:
            print(f"❌ Failed to update error status: {update_error}")


# ==================== PROPOSAL ENDPOINTS ====================

@app.post("/api/proposals")
async def create_proposal(
    proposal_type: str = Form(...),
    client_name: Optional[str] = Form(None),
    chat_file: Optional[UploadFile] = File(None),
    user=Depends(get_current_user)
):
    """
    Create a new proposal.
    Supports 'from_chat' with file upload.
    """
    db = get_database()
    if db is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Database not available"
        )
    
    user_id = user.get("user_id")
    
    # Validate proposal type
    valid_types = ["from_chat", "from_text"]
    if proposal_type not in valid_types:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid proposal_type. Must be one of: {', '.join(valid_types)}"
        )
    
    # Validate from_chat requirements
    if proposal_type == "from_chat":
        if not client_name:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="client_name is required for from_chat"
            )
        if not chat_file:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="chat_file is required for from_chat"
            )
        # Check file type (no PDFs, no old .doc)
        if chat_file.filename:
            file_ext = chat_file.filename.split('.')[-1].lower()
            if file_ext == 'pdf':
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="PDF files are not supported. Please upload a text document."
                )
            if file_ext == 'doc':
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Old .doc files are not supported. Please convert to .docx or upload .txt/.csv."
                )
    
    # Read file content if provided
    file_content = None
    file_name = None
    file_size = 0
    text_content = None
    
    if chat_file:
        file_content = await chat_file.read()
        file_name = chat_file.filename
        file_size = len(file_content)
        file_ext = file_name.split('.')[-1].lower()
        
        if file_ext == 'docx':
            try:
                doc = Document(io.BytesIO(file_content))
                text_content = "\n".join([p.text for p in doc.paragraphs])
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_content += "\n" + cell.text
            except Exception as e:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail=f"Failed to parse DOCX file: {e}"
                )
        elif file_ext == 'csv':
            try:
                decoded_content = file_content.decode('utf-8')
                reader = csv.reader(io.StringIO(decoded_content))
                text_content = "\n".join([",".join(row) for row in reader])
            except Exception as e:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail=f"Failed to parse CSV file: {e}"
                )
        else:  # Assume text file (.txt) or other plain text
            try:
                text_content = file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    text_content = file_content.decode('latin-1')
                except UnicodeDecodeError:
                    text_content = None
                    print(f"⚠️ Could not decode file {file_name} as UTF-8 or latin-1.")
        
        input_data = {
            "file_name": file_name,
            "file_size": file_size,
            "file_type": chat_file.content_type or "unknown",
            "chat_content": text_content,
            "has_full_content": text_content is not None
        }
    else:
        input_data = {}
    
    # Create proposal document
    proposal_doc = {
        "user_id": user_id,
        "proposal_type": proposal_type,
        "client_name": client_name,
        "status": "pending",
        "input_data": input_data,
        "results": None,
        "created_at": datetime.utcnow(),
        "updated_at": datetime.utcnow()
    }
    
    # Insert into database
    collection = db["proposals"]
    result = await collection.insert_one(proposal_doc)
    proposal_id = result.inserted_id
    
    # Start proposal generation in background (for from_chat)
    if proposal_type == "from_chat" and text_content:
        # Update status to processing
        await collection.update_one(
            {"_id": proposal_id},
            {"$set": {"status": "processing", "updated_at": datetime.utcnow()}}
        )
        
        # Run proposal generation in background (using asyncio.create_task)
        import asyncio
        asyncio.create_task(process_proposal(str(proposal_id), text_content, collection))
    
    # Return created proposal
    created_doc = await collection.find_one({"_id": proposal_id})
    created_doc["_id"] = str(created_doc["_id"])
    
    return created_doc


@app.get("/api/proposals")
async def get_proposals(
    user=Depends(get_current_user),
    limit: int = 50,
    skip: int = 0
):
    """Get all proposals for the current user."""
    db = get_database()
    if db is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Database not available"
        )
    
    user_id = user.get("user_id")
    collection = db["proposals"]
    
    # Query user's proposals, sorted by created_at descending
    cursor = collection.find({"user_id": user_id}).sort("created_at", -1).skip(skip).limit(limit)
    proposals = await cursor.to_list(length=limit)
    
    # Convert ObjectId to string
    for proposal in proposals:
        proposal["_id"] = str(proposal["_id"])
    
    return {"proposals": proposals, "count": len(proposals)}


@app.get("/api/proposals/{proposal_id}")
async def get_proposal(
    proposal_id: str,
    user=Depends(get_current_user)
):
    """Get a specific proposal by ID."""
    db = get_database()
    if db is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Database not available"
        )
    
    user_id = user.get("user_id")
    collection = db["proposals"]
    
    try:
        proposal = await collection.find_one({
            "_id": ObjectId(proposal_id),
            "user_id": user_id
        })
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid proposal ID"
        )
    
    if not proposal:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Proposal not found"
        )
    
    proposal["_id"] = str(proposal["_id"])
    return proposal


@app.delete("/api/proposals/{proposal_id}")
async def delete_proposal(
    proposal_id: str,
    user=Depends(get_current_user)
):
    """Delete a specific proposal by ID."""
    db = get_database()
    if db is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Database not available"
        )
    
    user_id = user.get("user_id")
    collection = db["proposals"]
    
    try:
        delete_result = await collection.delete_one({
            "_id": ObjectId(proposal_id),
            "user_id": user_id
        })
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid proposal ID"
        )
    
    if delete_result.deleted_count == 0:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Proposal not found or not owned by user"
        )
    
    return {"message": "Proposal deleted successfully"}


@app.post("/api/proposals/{proposal_id}/generate-risk-report")
async def generate_risk_report_from_proposal(
    proposal_id: str,
    user=Depends(get_current_user)
):
    """
    Generate a risk analysis report from an existing proposal.
    Uses the chat content from the proposal's input_data.
    """
    db = get_database()
    if db is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Database not available"
        )
    
    user_id = user.get("user_id")
    collection = db["proposals"]
    
    # Find the proposal
    try:
        proposal = await collection.find_one({
            "_id": ObjectId(proposal_id),
            "user_id": user_id
        })
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid proposal ID"
        )
    
    if not proposal:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Proposal not found or not owned by user"
        )
    
    # Check if proposal has chat content
    if not proposal.get("input_data") or not proposal.get("input_data", {}).get("chat_content"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Proposal does not have chat content. Cannot generate risk report."
        )
    
    chat_content = proposal["input_data"]["chat_content"]
    client_name = proposal.get("client_name") or "Untitled Analysis"
    
    # Create risk analysis document
    risk_collection = db["risk_analyses"]
    risk_doc = {
        "user_id": user_id,
        "analysis_type": "client_chat_import",
        "client_name": client_name,
        "status": "pending",
        "input_data": {
            "file_name": proposal["input_data"].get("file_name", "from_proposal"),
            "file_size": proposal["input_data"].get("file_size", 0),
            "file_type": proposal["input_data"].get("file_type", "unknown"),
            "chat_content": chat_content,
            "has_full_content": True
        },
        "results": None,
        "created_at": datetime.utcnow(),
        "updated_at": datetime.utcnow()
    }
    
    # Insert into database
    result = await risk_collection.insert_one(risk_doc)
    risk_analysis_id = result.inserted_id
    
    # Update status to processing
    await risk_collection.update_one(
        {"_id": risk_analysis_id},
        {"$set": {"status": "processing", "updated_at": datetime.utcnow()}}
    )
    
    # Start risk analysis in background
    import asyncio
    asyncio.create_task(process_risk_analysis(str(risk_analysis_id), chat_content, risk_collection))
    
    # Return created risk analysis
    created_doc = await risk_collection.find_one({"_id": risk_analysis_id})
    created_doc["_id"] = str(created_doc["_id"])
    
    return created_doc


@app.get("/api/proposals/{proposal_id}/history")
async def get_proposal_history(
    proposal_id: str,
    user=Depends(get_current_user)
):
    """Get the version history for a proposal."""
    db = get_database()
    if db is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Database not available"
        )
    
    user_id = user.get("user_id")
    collection = db["proposals"]
    
    try:
        proposal = await collection.find_one({
            "_id": ObjectId(proposal_id),
            "user_id": user_id
        })
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid proposal ID"
        )
    
    if not proposal:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Proposal not found or not owned by user"
        )
    
    # Return history array and current version index
    history = proposal.get("history", [])
    current_version_index = proposal.get("current_version_index", -1)
    
    # Convert ObjectId and datetime to strings for JSON serialization
    for entry in history:
        if isinstance(entry.get("timestamp"), datetime):
            entry["timestamp"] = entry["timestamp"].isoformat()
    
    return {
        "history": history,
        "current_version_index": current_version_index,
        "total_versions": len(history)
    }


@app.post("/api/proposals/{proposal_id}/restore")
async def restore_proposal_version(
    proposal_id: str,
    version_index: int = Form(...),
    user=Depends(get_current_user)
):
    """Restore a proposal to a previous version."""
    db = get_database()
    if db is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Database not available"
        )
    
    user_id = user.get("user_id")
    collection = db["proposals"]
    
    try:
        proposal = await collection.find_one({
            "_id": ObjectId(proposal_id),
            "user_id": user_id
        })
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid proposal ID"
        )
    
    if not proposal:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Proposal not found or not owned by user"
        )
    
    history = proposal.get("history", [])
    
    if version_index < 0 or version_index >= len(history):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid version index"
        )
    
    # Get the version to restore from history
    version_to_restore = history[version_index]
    
    # Update the proposal with the restored version
    if proposal.get("results"):
        proposal["results"]["formatted_proposal"] = version_to_restore["formatted_proposal"]
        proposal["results"]["updated_at"] = datetime.utcnow().isoformat()
    else:
        proposal["results"] = {
            "formatted_proposal": version_to_restore["formatted_proposal"],
            "updated_at": datetime.utcnow().isoformat()
        }
    
    # Update current_version_index
    proposal["current_version_index"] = version_index
    
    await collection.update_one(
        {"_id": ObjectId(proposal_id)},
        {
            "$set": {
                "results": proposal["results"],
                "current_version_index": version_index,
                "updated_at": datetime.utcnow()
            }
        }
    )
    
    # Return updated proposal
    updated_doc = await collection.find_one({"_id": ObjectId(proposal_id)})
    updated_doc["_id"] = str(updated_doc["_id"])
    
    return updated_doc


@app.patch("/api/proposals/{proposal_id}")
async def update_proposal_endpoint(
    proposal_id: str,
    user_changes: str = Form(...),
    chat_file: Optional[UploadFile] = File(None),
    user=Depends(get_current_user)
):
    """
    Update an existing proposal based on user-specified changes.
    Optionally accepts a new chat file to incorporate additional information.
    """
    db = get_database()
    if db is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Database not available"
        )
    
    user_id = user.get("user_id")
    collection = db["proposals"]
    
    # Find the proposal
    try:
        proposal = await collection.find_one({
            "_id": ObjectId(proposal_id),
            "user_id": user_id
        })
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid proposal ID"
        )
    
    if not proposal:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Proposal not found or not owned by user"
        )
    
    # Check if proposal has results
    if not proposal.get("results") or not proposal.get("results", {}).get("formatted_proposal"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Proposal has not been generated yet. Cannot update."
        )
    
    current_proposal = proposal["results"]["formatted_proposal"]
    
    # Parse new chat file if provided
    new_chat_content = None
    if chat_file:
        file_content = await chat_file.read()
        file_name = chat_file.filename
        file_ext = file_name.split('.')[-1].lower() if file_name else ''
        
        # Validate file type
        if file_ext == 'pdf':
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="PDF files are not supported. Please upload a text document (.txt, .docx, .csv)."
            )
        if file_ext == 'doc':
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Old .doc files are not supported. Please convert to .docx or upload .txt/.csv."
            )
        
        # Parse file content
        if file_ext == 'docx':
            try:
                doc = Document(io.BytesIO(file_content))
                new_chat_content = "\n".join([p.text for p in doc.paragraphs])
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            new_chat_content += "\n" + cell.text
            except Exception as e:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail=f"Failed to parse DOCX file: {e}"
                )
        elif file_ext == 'csv':
            try:
                decoded_content = file_content.decode('utf-8')
                reader = csv.reader(io.StringIO(decoded_content))
                new_chat_content = "\n".join([",".join(row) for row in reader])
            except Exception as e:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail=f"Failed to parse CSV file: {e}"
                )
        else:  # Assume text file (.txt) or other plain text
            try:
                new_chat_content = file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    new_chat_content = file_content.decode('latin-1')
                except UnicodeDecodeError:
                    raise HTTPException(
                        status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                        detail=f"Could not decode file {file_name} as UTF-8 or latin-1."
                    )
    
    # Update status to processing
    await collection.update_one(
        {"_id": ObjectId(proposal_id)},
        {"$set": {"status": "processing", "updated_at": datetime.utcnow()}}
    )
    
    # Process update in background
    import asyncio
    asyncio.create_task(process_proposal_update(
        str(proposal_id), 
        current_proposal, 
        user_changes, 
        new_chat_content,
        collection
    ))
    
    # Return updated proposal document (status will be processing)
    updated_doc = await collection.find_one({"_id": ObjectId(proposal_id)})
    updated_doc["_id"] = str(updated_doc["_id"])
    
    return updated_doc


@app.put("/api/proposals/{proposal_id}/save")
async def save_edited_proposal(
    proposal_id: str,
    formatted_proposal: str = Form(...),
    user=Depends(get_current_user)
):
    """
    Save a manually edited proposal and update history.
    """
    db = get_database()
    if db is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Database not available"
        )
    
    user_id = user.get("user_id")
    collection = db["proposals"]
    
    # Find the proposal
    try:
        proposal = await collection.find_one({
            "_id": ObjectId(proposal_id),
            "user_id": user_id
        })
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid proposal ID"
        )
    
    if not proposal:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Proposal not found or not owned by user"
        )
    
    # Check if proposal has results
    if not proposal.get("results") or not proposal.get("results", {}).get("formatted_proposal"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Proposal has not been generated yet. Cannot save."
        )
    
    # Initialize history if it doesn't exist
    if "history" not in proposal:
        proposal["history"] = []
    
    # Get current version index
    current_version_index = proposal.get("current_version_index", len(proposal["history"]) - 1)
    
    # Save current version to history before updating
    current_proposal_text = proposal["results"]["formatted_proposal"]
    
    # Only save to history if the content has actually changed
    if current_proposal_text != formatted_proposal:
        # If we're not at the latest version, truncate history (linear history only)
        if current_version_index < len(proposal["history"]) - 1:
            proposal["history"] = proposal["history"][:current_version_index + 1]
        
        # Add current version to history
        history_entry = {
            "formatted_proposal": current_proposal_text,
            "timestamp": datetime.utcnow(),
            "version": len(proposal["history"]) + 1
        }
        proposal["history"].append(history_entry)
        
        # Add the new edited version to history
        new_version_entry = {
            "formatted_proposal": formatted_proposal,
            "timestamp": datetime.utcnow(),
            "version": len(proposal["history"]) + 1
        }
        proposal["history"].append(new_version_entry)
        
        # Update current_version_index to point to the new version
        current_version_index = len(proposal["history"]) - 1
    
    # Update the proposal
    if proposal.get("results"):
        proposal["results"]["formatted_proposal"] = formatted_proposal
        proposal["results"]["updated_at"] = datetime.utcnow().isoformat()
    else:
        proposal["results"] = {
            "formatted_proposal": formatted_proposal,
            "updated_at": datetime.utcnow().isoformat()
        }
    
    await collection.update_one(
        {"_id": ObjectId(proposal_id)},
        {
            "$set": {
                "results": proposal["results"],
                "history": proposal["history"],
                "current_version_index": current_version_index,
                "updated_at": datetime.utcnow()
            }
        }
    )
    
    # Return updated proposal
    updated_doc = await collection.find_one({"_id": ObjectId(proposal_id)})
    updated_doc["_id"] = str(updated_doc["_id"])
    
    return updated_doc


async def process_proposal_update(
    proposal_id: str,
    current_proposal: str,
    user_changes: str,
    new_chat_content: Optional[str],
    collection
):
    """
    Background task to update a proposal using AI.
    """
    try:
        # Call AI service to update proposal
        updated_proposal_text = await update_proposal(
            current_proposal=current_proposal,
            user_changes=user_changes,
            new_chat_content=new_chat_content
        )
        
        # Get current proposal to save to history
        proposal = await collection.find_one({"_id": ObjectId(proposal_id)})
        
        # Initialize history array if it doesn't exist
        if "history" not in proposal:
            proposal["history"] = []
        
        # Save current version to history before updating
        current_version_index = proposal.get("current_version_index", len(proposal["history"]) - 1)
        if proposal.get("results") and proposal["results"].get("formatted_proposal"):
            # If we're not at the latest version, we need to truncate history
            # (can't have branches - only linear history)
            if current_version_index < len(proposal["history"]) - 1:
                proposal["history"] = proposal["history"][:current_version_index + 1]
            
            history_entry = {
                "formatted_proposal": proposal["results"]["formatted_proposal"],
                "timestamp": datetime.utcnow(),
                "version": len(proposal["history"]) + 1
            }
            proposal["history"].append(history_entry)
        
        # After updating, add the new version to history
        new_version_entry = {
            "formatted_proposal": updated_proposal_text,
            "timestamp": datetime.utcnow(),
            "version": len(proposal["history"]) + 1
        }
        proposal["history"].append(new_version_entry)
        
        # Set current_version_index to the new version (last in history)
        current_version_index = len(proposal["history"]) - 1
        
        # Update the proposal in database
        # We need to preserve the existing results structure but update formatted_proposal
        if proposal and proposal.get("results"):
            # Update only the formatted_proposal, preserve other fields
            proposal["results"]["formatted_proposal"] = updated_proposal_text
            proposal["results"]["updated_at"] = datetime.utcnow().isoformat()
        else:
            # If results don't exist, create them
            proposal["results"] = {
                "formatted_proposal": updated_proposal_text,
                "updated_at": datetime.utcnow().isoformat()
            }
        
        await collection.update_one(
            {"_id": ObjectId(proposal_id)},
            {
                "$set": {
                    "status": "completed",
                    "results": proposal["results"],
                    "history": proposal["history"],
                    "current_version_index": current_version_index,
                    "updated_at": datetime.utcnow()
                }
            }
        )
        
        print(f"✅ Proposal {proposal_id} updated successfully")
        
    except Exception as e:
        print(f"❌ Error updating proposal {proposal_id}: {e}")
        # Update status to failed
        await collection.update_one(
            {"_id": ObjectId(proposal_id)},
            {
                "$set": {
                    "status": "failed",
                    "updated_at": datetime.utcnow()
                }
            }
        )

